"""
==========================================================================
BACKEND TỔNG HỢP: TÁI CẤU TRÚC VĂN BẢN PHÁP LUẬT VIỆT NAM
==========================================================================
Pipeline:
    1. Ảnh chụp → PaddleOCR (detect box + OCR thô)
    2. → Sửa lỗi tiếng Việt bằng AI (Seq2Seq)
    3. → Hậu xử lý Dictionary (box ngắn < 20 ký tự)
    4. → Xác định dòng Title (nếu là trang đầu)
    5. → Xác định đoạn Title dài (extended title)
    6. → Chia vùng Header / Body
    7. → Gom đoạn Body (4 tiêu chí)
    8. → Xuất file Word (.docx) theo Nghị định 30/2020/NĐ-CP

Format (NĐ 30/2020/NĐ-CP):
    Lề trên: 20mm, Lề dưới: 20mm, Lề trái: 30mm, Lề phải: 18mm
    Font: Times New Roman 13pt
    Table header: line spacing 1.0, space_before=0, space_after=0
    Body: line spacing 1.0, space_before=0, space_after=6pt
    Title + ext_title: line spacing 1.0, space_before=0, space_after=6pt
    Enter trước title / sau ext_title: space_before=6pt, space_after=6pt
==========================================================================
"""

import cv2
import numpy as np
import torch
import unicodedata
import difflib
import re
from collections import Counter
from paddleocr import PaddleOCR
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ==========================================================================
# PHẦN 1: KHỞI TẠO MODEL (CHỈ CHẠY 1 LẦN DUY NHẤT)
# ==========================================================================

MODEL_DIR = r"D:\2_BK\NAM_4\test-for-LVTN\my_model\checkpoint-34125"
MODEL_PREFIX = "correction: "
MAX_INPUT_LENGTH = 300

print("[INIT] Đang khởi tạo PaddleOCR...")
ocr = PaddleOCR(use_angle_cls=True, lang='vi', show_log=False)

print(f"[INIT] Đang tải Model NLP từ: {MODEL_DIR}")
device = "cuda" if torch.cuda.is_available() else "cpu"
try:
    tokenizer = AutoTokenizer.from_pretrained(MODEL_DIR, use_fast=False)
    nlp_model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_DIR).to(device)
    nlp_model.eval()
    print(f"[INIT] ✅ Model NLP sẵn sàng trên {device}")
except Exception as e:
    print(f"[INIT] ❌ LỖI tải Model NLP: {e}")
    exit()


# ==========================================================================
# PHẦN 2: BỘ TỪ ĐIỂN + REGEX
# ==========================================================================

TITLE_DICT = [
    "HIẾN PHÁP", "LUẬT", "BỘ LUẬT", "NGHỊ QUYẾT", "NGHỊ ĐỊNH",
    "QUYẾT ĐỊNH", "THÔNG TƯ", "THÔNG TƯ LIÊN TỊCH", "CHỈ THỊ",
    "PHÁP LỆNH", "NGHỊ QUYẾT LIÊN TỊCH", "LỆNH", "QUY CHẾ", "QUY ĐỊNH"
]

AGENCY_DICT = [
    "QUỐC HỘI", "CHỦ TỊCH NƯỚC", "CHÍNH PHỦ", "THỦ TƯỚNG CHÍNH PHỦ",
    "TÒA ÁN NHÂN DÂN TỐI CAO", "VIỆN KIỂM SÁT NHÂN DÂN TỐI CAO",
    "KIỂM TOÁN NHÀ NƯỚC", "BỘ TƯ PHÁP", "BỘ CÔNG AN", "BỘ QUỐC PHÒNG",
    "BỘ TÀI CHÍNH", "BỘ NỘI VỤ", "BỘ NGOẠI GIAO",
    "BỘ NÔNG NGHIỆP VÀ PHÁT TRIỂN NÔNG THÔN", "BỘ GIAO THÔNG VẬN TẢI",
    "BỘ XÂY DỰNG", "BỘ TÀI NGUYÊN VÀ MÔI TRƯỜNG",
    "BỘ THÔNG TIN VÀ TRUYỀN THÔNG", "BỘ LAO ĐỘNG - THƯƠNG BINH VÀ XÃ HỘI",
    "BỘ VĂN HÓA, THỂ THAO VÀ DU LỊCH", "BỘ KHOA HỌC VÀ CÔNG NGHỆ",
    "BỘ GIÁO DỤC VÀ ĐÀO TẠO", "BỘ Y TẾ", "BỘ KẾ HOẠCH VÀ ĐẦU TƯ",
    "BỘ CÔNG THƯƠNG", "NGÂN HÀNG NHÀ NƯỚC VIỆT NAM",
    "ỦY BAN DÂN TỘC", "HỘI ĐỒNG NHÂN DÂN", "ỦY BAN NHÂN DÂN",
    "ỦY BAN THƯỜNG VỤ QUỐC HỘI"
]

PARAGRAPH_BREAK_KEYWORDS = re.compile(
    r"^\s*[\"""\'\'\']?\s*"
    r"(Điều\s*\d+[a-z]?\."
    r"|Chương\s+[IVXLCDM]+"
    r"|Bổ sung\b|Sửa đổi\b|Căn cứ\b|Theo đề nghị\b"
    r"|\d+[a-zđ]?\."           # khoản 1. 2. 3.
    r"|[a-zđ]\)"               # ← thêm đ vào đây
    r")",
    re.IGNORECASE
)

BOLD_KEYWORDS = re.compile(
    r"^\s*[\"""\'\'\']?\s*"
    r"(Điều\s*\d+[a-z]?\."
    r"|Chương\s+[IVXLCDM]+"
    r"|\d+[a-z]?\.)"
    , re.IGNORECASE
)


# ==========================================================================
# PHẦN 3: HÀM PHÁT HIỆN HALLUCINATION
# ==========================================================================

def is_hallucinated(text):
    if len(text) < 30:
        return False
    words = text.split()
    if len(words) < 5:
        return False
    single_char_ratio = sum(1 for w in words if len(w) == 1) / len(words)
    if single_char_ratio > 0.40:
        return True
    repeat_pattern = re.search(r'((?:\S+\s+){1,3}?)\1{4,}', text)
    if repeat_pattern:
        return True
    if len(words) > 20:
        unique_ratio = len(set(w.lower() for w in words)) / len(words)
        if unique_ratio < 0.15:
            return True
    return False


# ==========================================================================
# PHẦN 4: CÁC HÀM TIỆN ÍCH
# ==========================================================================

def correct_text_ai(raw_text):
    if not raw_text.strip():
        return ""
    inputs = tokenizer(
        MODEL_PREFIX + raw_text, return_tensors="pt",
        max_length=512, truncation=True
    ).to(device)

    num_beams = get_num_beams(raw_text)          # ← thêm dòng này

    with torch.no_grad():
        outputs = nlp_model.generate(
            inputs["input_ids"],
            max_length=512,
            num_beams=num_beams,                 # ← thay hardcode 5
            early_stopping=(num_beams > 1),      # ← tắt khi greedy
        )
    corrected = tokenizer.decode(outputs[0], skip_special_tokens=True).strip()
    if is_hallucinated(corrected):
        return raw_text
    return corrected


def chunk_paragraph(text, max_length=MAX_INPUT_LENGTH):
    if len(text) <= max_length:
        return [text]
    chunks = []
    while len(text) > 0:
        if len(text) <= max_length:
            chunks.append(text)
            break
        cut_idx = -1
        for punct in ['. ', '; ', '? ', '! ', '.', ';', '?', '!']:
            idx = text.rfind(punct, 0, max_length)
            if idx != -1:
                cut_idx = idx + len(punct)
                break
        if cut_idx == -1:
            for punct in [', ', ',']:
                idx = text.rfind(punct, 0, max_length)
                if idx != -1:
                    cut_idx = idx + len(punct)
                    break
        if cut_idx == -1:
            idx = text.rfind(' ', 0, max_length)
            cut_idx = idx if idx != -1 else max_length
        chunk = text[:cut_idx].strip()
        if chunk:
            chunks.append(chunk)
        text = text[cut_idx:].strip()
    return chunks


def correct_chunk(raw_chunk):
    if not raw_chunk.strip():
        return ""
    inputs = tokenizer(
        MODEL_PREFIX + raw_chunk, return_tensors="pt",
        max_length=512, truncation=True
    ).to(device)

    num_beams = get_num_beams(raw_chunk)         # ← thêm dòng này

    with torch.no_grad():
        outputs = nlp_model.generate(
            inputs["input_ids"],
            max_length=512,
            num_beams=num_beams,                 # ← thay hardcode 5
            no_repeat_ngram_size=3,
            early_stopping=(num_beams > 1),      # ← tắt khi greedy
        )
    corrected = tokenizer.decode(outputs[0], skip_special_tokens=True).strip()
    if is_hallucinated(corrected):
        return raw_chunk
    return corrected

# ==========================================================================
# HÀM BATCH INFERENCE (MỚI)
# Gom nhiều text → tokenize padding → 1 lần generate → decode từng output
# ==========================================================================

def correct_texts_batch(raw_texts: list[str], batch_size: int = 16) -> list[str]:
    """
    Thay thế việc gọi correct_text_ai() từng cái một.
    
    Args:
        raw_texts: danh sách text thô (có thể rỗng)
        batch_size: số sample mỗi lần forward (giảm nếu OOM)
    Returns:
        danh sách text đã sửa, cùng thứ tự với input
    """
    if not raw_texts:
        return []

    results = [""] * len(raw_texts)

    # Tách index của các text không rỗng để xử lý
    non_empty_indices = [i for i, t in enumerate(raw_texts) if t.strip()]
    non_empty_texts   = [raw_texts[i] for i in non_empty_indices]

    for batch_start in range(0, len(non_empty_texts), batch_size):
        batch_texts = non_empty_texts[batch_start : batch_start + batch_size]
        prefixed    = [MODEL_PREFIX + t for t in batch_texts]

        # Tokenize với padding → tensor shape [B, seq_len]
        inputs = tokenizer(
            prefixed,
            return_tensors="pt",
            max_length=512,
            truncation=True,
            padding=True,          # ← padding để stack thành batch
        ).to(device)

        # Tính num_beams cho cả batch (lấy max để đảm bảo chất lượng)
        # Hoặc dùng greedy (num_beams=1) để tối đa tốc độ
        max_len = max(len(t) for t in batch_texts)
        num_beams = get_num_beams_batch(max_len)

        with torch.no_grad():
            outputs = nlp_model.generate(
                inputs["input_ids"],
                attention_mask=inputs["attention_mask"],
                max_length=512,
                num_beams=num_beams,
                no_repeat_ngram_size=3 if num_beams > 1 else 0,
                early_stopping=(num_beams > 1),
            )

        # Decode từng output trong batch
        for local_idx, (out, raw) in enumerate(zip(outputs, batch_texts)):
            corrected = tokenizer.decode(out, skip_special_tokens=True).strip()
            if is_hallucinated(corrected):
                corrected = raw
            global_idx = non_empty_indices[batch_start + local_idx]
            results[global_idx] = corrected

    return results

# new
def get_num_beams_batch(max_text_len: int) -> int:
    """
    Phiên bản cho batch: dùng max độ dài trong batch để quyết định beam.
    Ưu tiên greedy/beam nhỏ để batch nhanh hơn.
    """
    if max_text_len <= 10:
        return 1
    elif max_text_len <= 30:
        return 2
    elif max_text_len <= 80:
        return 3
    else:
        return 5

def fuzzy_match_dict(text, dictionary, threshold=0.85):
    """So khớp mờ text với dictionary.
    Thêm logic: so sánh cả bản bỏ khoảng trắng để bắt lỗi dính chữ.
    VD: "PHÁPLENH" vs "PHÁP LỆNH" → bỏ space: "PHÁPLENH" vs "PHÁPLỆNH" → match
    """
    text_upper = text.upper().strip()
    text_no_space = text_upper.replace(" ", "")
    best_match, highest_ratio = None, 0.0

    for valid_word in dictionary:
        # So sánh bình thường (có space)
        ratio1 = difflib.SequenceMatcher(None, text_upper, valid_word).ratio()

        # So sánh bỏ space (bắt lỗi dính chữ)
        valid_no_space = valid_word.replace(" ", "")
        ratio2 = difflib.SequenceMatcher(None, text_no_space, valid_no_space).ratio()

        # Lấy ratio cao nhất
        ratio = max(ratio1, ratio2)

        if ratio > highest_ratio:
            highest_ratio = ratio
            best_match = valid_word

    if highest_ratio >= threshold:
        return best_match, highest_ratio
    return None, 0.0


def should_be_bold(text):
    text_lower = text.lower()
    if "/" in text or "số:" in text_lower or "s6:" in text_lower:
        return False
    if "ngày" in text_lower and "tháng" in text_lower:
        return False
    return True


def hardcode_quoc_hieu(text):
    text_upper = text.upper().strip()
    target = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    if "CỘNG" in text_upper and "VIỆT" in text_upper:
        ratio = difflib.SequenceMatcher(None, text_upper, target).ratio()
        if ratio >= 0.75:
            return target
    ratio = difflib.SequenceMatcher(None, text_upper, target).ratio()
    if ratio >= 0.80:
        return target
    return text


# ==========================================================================
# PHẦN 5: OCR + SỬA LỖI AI + HẬU XỬ LÝ DICTIONARY
# ==========================================================================

def run_ocr_and_postprocess(img, progress_cb=None):
    img_height, img_width = img.shape[:2]

    if progress_cb:
        progress_cb(0, 1, "Đang quét PaddleOCR...")

    raw_result = ocr.ocr(img, cls=True)[0]
    if not raw_result:
        return [], 0, img_height, img_width

    heights = [max(b[0][2][1], b[0][3][1]) - min(b[0][0][1], b[0][1][1])
               for b in raw_result]
    avg_char_height = np.median(heights)

    # ── [CŨ] Gọi AI từng box ──────────────────────────────────────────────
    # for idx, box_data in enumerate(raw_result):
    #     corrected = correct_text_ai(raw_text)  ← XÓA

    # ── [MỚI] Thu thập raw_text toàn bộ → batch inference 1 lần ──────────
    if progress_cb:
        progress_cb(0, 1, f"Đang sửa lỗi AI batch {len(raw_result)} box...")

    raw_texts = [box_data[1][0] for box_data in raw_result]
    corrected_texts = correct_texts_batch(raw_texts)   # ← 1 LẦN DUY NHẤT

    boxes = []
    for idx, (box_data, corrected) in enumerate(zip(raw_result, corrected_texts)):
        corrected = unicodedata.normalize('NFC', corrected)
        coords = box_data[0]
        x_min = min(pt[0] for pt in coords)
        x_max = max(pt[0] for pt in coords)
        y_min = min(pt[1] for pt in coords)
        y_max = max(pt[1] for pt in coords)
        boxes.append({
            'text': corrected, 'coords': coords,
            'x_min': x_min, 'x_max': x_max,
            'y_min': y_min, 'y_max': y_max,
            'y_center': (y_min + y_max) / 2,
            'x_center': (x_min + x_max) / 2,
        })

    for box in boxes:
        if len(box['text']) < 20:
            matched, ratio = fuzzy_match_dict(box['text'], TITLE_DICT, threshold=0.85)
            if matched:
                box['text'] = matched

    if progress_cb:
        progress_cb(1, 1, "OCR + AI batch + Dictionary hoàn tất!")

    return boxes, avg_char_height, img_height, img_width


# ==========================================================================
# PHẦN 6: GOM DÒNG
# ==========================================================================

def group_boxes_into_lines(boxes, avg_char_height):
    if not boxes:
        return []
    sorted_boxes = sorted(boxes, key=lambda b: b['y_center'])
    lines = []
    current_line = [sorted_boxes[0]]
    for box in sorted_boxes[1:]:
        prev_y = sum(b['y_center'] for b in current_line) / len(current_line)
        if abs(box['y_center'] - prev_y) < avg_char_height * 0.5:
            current_line.append(box)
        else:
            lines.append(current_line)
            current_line = [box]
    if current_line:
        lines.append(current_line)
    for line in lines:
        line.sort(key=lambda b: b['x_min'])
    return lines


# ==========================================================================
# PHẦN 7: TÌM DÒNG TITLE + ĐOẠN TITLE DÀI
# ==========================================================================

def find_title_line_index(lines, img_height, img_width):
    page_center_x = img_width / 2
    page_center_y = img_height / 2

    for i, line in enumerate(lines):
        line_text = " ".join([b['text'] for b in line]).strip()
        if len(line_text) >= 19:
            continue
        if line_text.upper() in TITLE_DICT:
            print(f"[TITLE] ✅ Match 100%: '{line_text}' tại dòng {i}")
            return i

    best_idx, best_score = -1, float('inf')
    for i, line in enumerate(lines):
        line_text = " ".join([b['text'] for b in line]).strip()
        text_len = len(line_text)
        if text_len > 25 or text_len < 2:
            continue
        line_cx = sum(b['x_center'] for b in line) / len(line)
        line_cy = sum(b['y_center'] for b in line) / len(line)
        dist = np.sqrt((line_cx - page_center_x)**2 + (line_cy - page_center_y)**2)
        diag = np.sqrt(img_width**2 + img_height**2)
        score = text_len + (dist / diag) * 50
        if score < best_score:
            best_score, best_idx = score, i
    return best_idx


def find_extended_title_range(lines, title_line_idx):
    if title_line_idx < 0 or title_line_idx >= len(lines) - 1:
        return (title_line_idx + 1, title_line_idx + 1)

    ext_start = title_line_idx + 1
    if ext_start >= len(lines):
        return (ext_start, ext_start)

    gaps = []
    for i in range(ext_start, min(len(lines) - 1, ext_start + 15)):
        y_max_prev = max(b['y_max'] for b in lines[i])
        y_min_next = min(b['y_min'] for b in lines[i + 1])
        gap = y_min_next - y_max_prev
        if gap > 0:
            gaps.append(gap)

    if not gaps:
        return (ext_start, ext_start + 1)

    normal_gap = np.median(gaps)
    enter_threshold = normal_gap * 1.5

    first_text = " ".join([b['text'] for b in lines[ext_start]]).strip()
    if first_text.lower().startswith("căn cứ"):
        return (ext_start, ext_start)

    ext_end = ext_start + 1

    for i in range(ext_start, len(lines) - 1):
        next_idx = i + 1
        next_text = " ".join([b['text'] for b in lines[next_idx]]).strip()
        if next_text.lower().startswith("căn cứ"):
            break
        y_max_curr = max(b['y_max'] for b in lines[i])
        y_min_next = min(b['y_min'] for b in lines[next_idx])
        actual_gap = y_min_next - y_max_curr
        if actual_gap <= enter_threshold:
            ext_end = next_idx + 1
        else:
            break

    return (ext_start, ext_end)


# ==========================================================================
# PHẦN 8: XUẤT HEADER RA WORD
# ĐÃ SỬA: Agency logic - so sánh 75% rồi bỏ box dưới nếu giống 85%
#          Table: line spacing 1.0, space_before=0, space_after=0
# ==========================================================================

def process_agency_lines(header_lines, avg_char_height, img_width):
    """
    Xử lý agency (góc trên trái):
        1. Gộp text cột trái từ nhiều dòng liên tiếp
        2. So sánh text gộp với AGENCY_DICT >= 75% → thay thế
        3. Các dòng cột trái đã bị gộp → bỏ text (giữ cột phải nếu có)
    """
    if not header_lines:
        return header_lines, None

    left_threshold = img_width * 0.4

    # Bước 1: Thu thập tất cả text cột trái từ các dòng liên tiếp (trước dòng Số hiệu)
    left_line_indices = []  # index của các dòng có box cột trái
    left_texts = []         # text cột trái tương ứng

    for i, line in enumerate(header_lines):
        left_boxes = [b for b in line if b['x_center'] < left_threshold]
        if not left_boxes:
            continue
        left_text = " ".join([b['text'] for b in left_boxes]).strip()
        # Dừng nếu gặp dòng Số hiệu (có "/")
        if "/" in left_text or "số:" in left_text.lower() or "s6:" in left_text.lower():
            break
        left_line_indices.append(i)
        left_texts.append(left_text)

    if not left_texts:
        return header_lines, None

    # Bước 2: Thử gộp text từ 1 dòng, 2 dòng, 3 dòng... rồi so sánh
    matched_agency = None
    matched_count = 0  # Số dòng cột trái được gộp

    for n in range(len(left_texts), 0, -1):
        # Gộp n dòng đầu tiên
        combined = " ".join(left_texts[:n]).strip()
        matched, ratio = fuzzy_match_dict(combined, AGENCY_DICT, threshold=0.75)
        if matched:
            matched_agency = matched
            matched_count = n
            print(f"  [AGENCY] Gộp {n} dòng: '{combined}' → '{matched}' ({ratio*100:.0f}%)")
            break

    if matched_agency is None:
        return header_lines, None

    # Bước 3: Xử lý các dòng
    # - Dòng đầu tiên cột trái: thay text thành matched_agency
    # - Các dòng cột trái còn lại (đã bị gộp): bỏ text cột trái, giữ cột phải
    lines_to_keep = []
    matched_line_indices = set(left_line_indices[:matched_count])

    for i, line in enumerate(header_lines):
        if i not in matched_line_indices:
            lines_to_keep.append(line)
            continue

        if i == left_line_indices[0]:
            # Dòng đầu: thay text cột trái thành matched_agency
            left_boxes = [b for b in line if b['x_center'] < left_threshold]
            right_boxes = [b for b in line if b['x_center'] >= left_threshold]
            for b in left_boxes:
                b['text'] = ""
            if left_boxes:
                left_boxes[0]['text'] = matched_agency
            lines_to_keep.append(line)
        else:
            # Các dòng sau: bỏ cột trái, chỉ giữ cột phải
            right_boxes = [b for b in line if b['x_center'] >= left_threshold]
            if right_boxes:
                lines_to_keep.append(right_boxes)
            else:
                print(f"  [AGENCY] Bỏ hoàn toàn dòng {i}")
                # Không thêm vào lines_to_keep → bỏ dòng này

    return lines_to_keep, matched_agency


def write_header_to_doc(doc, header_lines, avg_char_height, img_width):
    """
    Header → bảng 2 cột (tỉ lệ 3.5/6.5).
    Table: line spacing 1.0, space_before=0, space_after=0
    """
    # Xử lý agency trước
    header_lines, _ = process_agency_lines(header_lines, avg_char_height, img_width)

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False

    # Tỉ lệ cột 3.5/6.5
    COL_LEFT_WIDTH = Cm(5.67)
    COL_RIGHT_WIDTH = Cm(10.53)

    in_table_zone = True
    doc_lap_row_added = False

    for line in header_lines:
        full_text = " ".join([b['text'] for b in line])
        full_upper = full_text.upper()

        idx_cong = full_upper.find("CỘNG")
        idx_doc_lap = full_upper.find("ĐỘC")
        left_text, right_text, is_two_col = "", "", False

        if idx_cong != -1 and "VIỆT" in full_upper:
            left_text = full_text[:idx_cong].strip()
            right_text = full_text[idx_cong:].strip()
            is_two_col = True
        elif idx_doc_lap != -1 and ("LẬP" in full_upper or "TỰ" in full_upper):
            left_text = full_text[:idx_doc_lap].strip()
            right_text = full_text[idx_doc_lap:].strip()
            is_two_col = True
        elif len(line) >= 2 and in_table_zone:
            max_gap, split_idx = 0, 1
            for i in range(1, len(line)):
                gap = line[i]['x_min'] - line[i-1]['x_max']
                if gap > max_gap:
                    max_gap, split_idx = gap, i
            if max_gap > avg_char_height * 1.5:
                left_text = " ".join([b['text'] for b in line[:split_idx]])
                right_text = " ".join([b['text'] for b in line[split_idx:]])
                is_two_col = True

        if is_two_col:
            is_doc_lap_line = ("ĐỘC" in right_text.upper() and
                               ("LẬP" in right_text.upper() or "TỰ" in right_text.upper()))

            row = table.add_row()
            row.cells[0].width = COL_LEFT_WIDTH
            row.cells[1].width = COL_RIGHT_WIDTH

            # Cột trái: fuzzy AGENCY_DICT 65% (cho các dòng chưa xử lý ở process_agency)
            if not ("/" in left_text or "số:" in left_text.lower() or "s6:" in left_text.lower()):
                matched, ratio = fuzzy_match_dict(left_text, AGENCY_DICT, threshold=0.65)
                if matched:
                    left_text = matched

            p_left = row.cells[0].paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_left.paragraph_format.line_spacing = 1.0
            p_left.paragraph_format.space_after = Pt(0)
            p_left.paragraph_format.space_before = Pt(0)
            run_l = p_left.add_run(left_text)
            run_l.font.name = 'Times New Roman'
            run_l.font.size = Pt(13)
            run_l.bold = should_be_bold(left_text)

            # Cột phải: hard-code Quốc hiệu
            right_text = hardcode_quoc_hieu(right_text)
            if "ĐỘC" in right_text.upper() and ("LẬP" in right_text.upper() or "TỰ" in right_text.upper()):
                right_text = "Độc lập - Tự do - Hạnh phúc"

            p_right = row.cells[1].paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_right.paragraph_format.line_spacing = 1.0
            p_right.paragraph_format.space_after = Pt(0)
            p_right.paragraph_format.space_before = Pt(0)
            run_r = p_right.add_run(right_text)
            run_r.font.name = 'Times New Roman'
            run_r.font.size = Pt(13)
            run_r.bold = should_be_bold(right_text)
            if "ngày" in right_text.lower() or "tháng" in right_text.lower():
                run_r.italic = True

            # Dòng trắng sau "Độc lập..."
            if is_doc_lap_line and not doc_lap_row_added:
                empty_row = table.add_row()
                empty_row.cells[0].width = COL_LEFT_WIDTH
                empty_row.cells[1].width = COL_RIGHT_WIDTH
                for cell in empty_row.cells:
                    p = cell.paragraphs[0]
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                doc_lap_row_added = True

        else:
            in_table_zone = False
            if len(full_text) < 30:
                matched, _ = fuzzy_match_dict(full_text, TITLE_DICT, threshold=0.80)
                if matched:
                    full_text = matched
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(full_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = should_be_bold(full_text)


# ==========================================================================
# PHẦN 9: XUẤT ĐOẠN TITLE DÀI RA WORD
# ==========================================================================

def balance_extended_title(full_text, max_line_chars=60):
    """
    Cân bằng đoạn title dài cho đẹp mắt.

    max_line_chars: ước lượng số ký tự tối đa 1 dòng khi căn giữa trong Word
                    (font 13, khổ A4 lề 30+18mm ≈ 60 ký tự)

    Thuật toán:
        1. Text ngắn (<= max_line_chars) → giữ nguyên 1 dòng
        2. Text 1 dòng dài (> 65% max nhưng < 2 dòng đầy) → chia 2 dòng đều
        3. Text nhiều dòng + dòng cuối <= 60% max → cân bằng 2 dòng cuối
        4. Dòng cuối > 60% → không cần xử lý

    Returns: text đã chèn '\\n' ở vị trí xuống dòng
    """
    words = full_text.split()
    if not words:
        return full_text

    total_len = len(full_text)

    # Trường hợp 1: Text ngắn → giữ nguyên
    if total_len <= max_line_chars:
        return full_text

    # Trường hợp 2: Text vừa phải (> 65% nhưng < 2 dòng đầy)
    if total_len <= max_line_chars * 2:
        return _split_balanced(full_text, 2)

    # Trường hợp 3+4: Text dài nhiều dòng
    num_lines = -(-total_len // max_line_chars)  # ceil division
    lines = _split_balanced(full_text, num_lines).split('\n')

    # Kiểm tra dòng cuối: nếu <= 60% max → cân bằng lại 2 dòng cuối
    if len(lines) >= 2:
        last_line = lines[-1]
        prev_line = lines[-2]
        last_ratio = len(last_line) / max_line_chars

        if last_ratio <= 0.60:
            merged = prev_line + " " + last_line
            new_two = _split_balanced(merged, 2)
            lines = lines[:-2] + new_two.split('\n')

    return '\n'.join(lines)


def _split_balanced(text, num_parts):
    """
    Chia text thành num_parts phần có độ dài tương đương nhau.
    Cắt tại khoảng trắng gần nhất với vị trí chia đều.
    """
    if num_parts <= 1:
        return text

    words = text.split()
    if len(words) <= num_parts:
        return text

    total_len = len(text)
    target_len = total_len / num_parts

    parts = []
    current_pos = 0

    for part_idx in range(num_parts - 1):
        ideal_cut = int(target_len * (part_idx + 1))

        left_space = text.rfind(' ', current_pos, ideal_cut + 1)
        right_space = text.find(' ', ideal_cut)

        if left_space == -1 and right_space == -1:
            break
        elif left_space == -1:
            best_cut = right_space
        elif right_space == -1:
            best_cut = left_space
        else:
            if abs(left_space - ideal_cut) <= abs(right_space - ideal_cut):
                best_cut = left_space
            else:
                best_cut = right_space

        parts.append(text[current_pos:best_cut].strip())
        current_pos = best_cut + 1

    remaining = text[current_pos:].strip()
    if remaining:
        parts.append(remaining)

    return '\n'.join(parts)


def write_extended_title_to_doc(doc, ext_title_lines):
    """Xuất đoạn Title dài: CĂN GIỮA + BÔI ĐẬM, font 13, cân bằng dòng."""
    if not ext_title_lines:
        return
    full_text = " ".join(
        " ".join([b['text'] for b in line]).strip()
        for line in ext_title_lines
    ).strip()

    # Cân bằng độ dài các dòng cho đẹp mắt
    balanced_text = balance_extended_title(full_text)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    # Xuất từng dòng (tách bởi \n)
    lines = balanced_text.split('\n')
    for i, line_text in enumerate(lines):
        run = p.add_run(line_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(15)
        run.bold = True
        # Xuống dòng (trừ dòng cuối)
        if i < len(lines) - 1:
            p.add_run('\n').font.size = Pt(15)


# ==========================================================================
# PHẦN 10: GOM ĐOẠN BODY
# ==========================================================================

def compute_line_stats(body_lines, img_width):
    line_stats = []
    for line in body_lines:
        x_min = min(b['x_min'] for b in line)
        x_max = max(b['x_max'] for b in line)
        y_min = min(b['y_min'] for b in line)
        y_max = max(b['y_max'] for b in line)
        w = x_max - x_min
        text = " ".join([b['text'] for b in line]).strip()
        line_stats.append({
            'x_min': x_min, 'x_max': x_max,
            'y_min': y_min, 'y_max': y_max,
            'y_center': (y_min + y_max) / 2,
            'width': w, 'text': text,
            'line_data': line
        })

    valid_widths = [s['width'] for s in line_stats if s['width'] > img_width * 0.3]
    if valid_widths:
        full_line_width = Counter(
            [round(w / 10) * 10 for w in valid_widths]
        ).most_common(1)[0][0]
    else:
        full_line_width = img_width * 0.8

    left_candidates = [s['x_min'] for s in line_stats if s['x_min'] < img_width * 0.3]
    if left_candidates:
        rounded = [round(x / 10) * 10 for x in left_candidates]
        global_x_min = Counter(rounded).most_common(1)[0][0]
    else:
        global_x_min = 0

    return line_stats, full_line_width, global_x_min


def compute_normal_line_gap(line_stats):
    if len(line_stats) < 2:
        return 0
    gaps = []
    for i in range(1, len(line_stats)):
        gap = line_stats[i]['y_min'] - line_stats[i - 1]['y_max']
        if gap > 0:
            gaps.append(gap)
    if not gaps:
        return 0
    return np.median(gaps)


def group_lines_into_paragraphs(line_stats, full_line_width, normal_gap, global_x_min=0, avg_char_height=20):
    if not line_stats:
        return []

    paragraphs = []
    current_para = [line_stats[0]]

    for i in range(1, len(line_stats)):
        current_stat = line_stats[i]
        prev_stat = line_stats[i - 1]
        should_cut = False

        if PARAGRAPH_BREAK_KEYWORDS.match(current_stat['text']):
            should_cut = True
        elif normal_gap > 0:
            actual_gap = current_stat['y_min'] - prev_stat['y_max']
            if actual_gap > normal_gap * 2.0:
                should_cut = True

        # MỚI — cắt khi prev ngắn HẲN *VÀ* curr có indent (= thực sự là đầu đoạn)
        if not should_cut:
            prev_short = prev_stat['width'] < (full_line_width * 0.75)
            curr_indented = current_stat['x_min'] > global_x_min + avg_char_height * 0.6
            if prev_short and curr_indented:
                should_cut = True

        if should_cut:
            paragraphs.append(current_para)
            current_para = [current_stat]
        else:
            current_para.append(current_stat)

    if current_para:
        paragraphs.append(current_para)

    return paragraphs


# ==========================================================================
# PHẦN 11: XUẤT BODY RA WORD
# ĐÃ SỬA: line spacing 1.0, space_before=0, space_after=6pt
# ==========================================================================

def write_body_to_doc(doc, body_lines, avg_char_height, img_width,
                      progress_cb=None, progress_offset=0, progress_total=1):
    if not body_lines:
        return

    line_stats, full_line_width, global_x_min = compute_line_stats(body_lines, img_width)
    normal_gap = compute_normal_line_gap(line_stats)
    paragraphs = group_lines_into_paragraphs(
        line_stats, full_line_width, normal_gap, global_x_min, avg_char_height
    )

    if not paragraphs:
        return

    # ── [MỚI] Chuẩn bị tất cả chunk TRƯỚC, batch 1 lần ──────────────────
    # Bước 1: tạo danh sách (para_idx, chunk_idx, raw_chunk)
    para_raw_texts   = []   # text thô mỗi para (chưa chunk nhỏ)
    para_x_mins      = []
    para_chunks_map  = []   # [(start_flat_idx, end_flat_idx), ...]
    all_flat_chunks  = []   # tất cả chunk flatten

    for para_stats in paragraphs:
        raw_text = " ".join([s['text'] for s in para_stats])
        para_x_min = para_stats[0]['x_min']
        chunks = chunk_paragraph(raw_text)
        
        start = len(all_flat_chunks)
        all_flat_chunks.extend(chunks)
        end = len(all_flat_chunks)
        
        para_raw_texts.append(raw_text)
        para_x_mins.append(para_x_min)
        para_chunks_map.append((start, end))

    # Bước 2: batch inference toàn bộ chunks 1 lần
    if progress_cb:
        progress_cb(progress_offset, progress_total,
                    f"Sửa lỗi AI batch {len(all_flat_chunks)} chunk body...")

    corrected_flat = correct_texts_batch(all_flat_chunks)  # ← 1 LẦN DUY NHẤT

    # Bước 3: ghép lại và xuất Word
    for idx, (para_stats, para_x_min, (s, e)) in enumerate(
        zip(paragraphs, para_x_mins, para_chunks_map)
    ):
        corrected_text = " ".join(corrected_flat[s:e]).strip()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

        if para_x_min > global_x_min + (avg_char_height * 0.8):
            p.paragraph_format.first_line_indent = Cm(1.27)

        bold_match = BOLD_KEYWORDS.match(corrected_text)
        if bold_match:
            bold_part = bold_match.group(0)
            run1 = p.add_run(bold_part)
            run1.bold = True
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(13)
            run2 = p.add_run(corrected_text[len(bold_part):])
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(13)
        else:
            run = p.add_run(corrected_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)


# ==========================================================================
# PHẦN 12: VẼ ẢNH DEBUG
# ==========================================================================

def _put_label_left(img, label, x1, y1, y2, color, font_scale=0.45, thickness=1):
    """
    Vẽ chữ label dọc theo cạnh TRÁI của bounding box (căn giữa theo chiều cao).
    Chữ được đặt ngay bên ngoài (bên trái) của box, không đè lên nội dung.
    """
    font = cv2.FONT_HERSHEY_SIMPLEX
    (tw, th), baseline = cv2.getTextSize(label, font, font_scale, thickness)
    box_h = y2 - y1
    # Căn giữa label theo chiều dọc của box
    ty = y1 + (box_h + th) // 2
    # Đặt chữ bên trái box (lùi ra ngoài một chút)
    tx = max(0, x1 - tw - 4)
    cv2.putText(img, label, (tx, ty), font, font_scale, color, thickness, cv2.LINE_AA)


def draw_debug_image(img, boxes, lines, title_line_idx,
                     ext_title_range, body_lines,
                     avg_char_height, img_width):
    debug_img = img.copy()

    LABEL_FONT_SCALE = 0.45
    LABEL_THICKNESS  = 1

    # ------------------------------------------------------------------
    # 1. Xác định các box thuộc TITLE, HEADER
    # ------------------------------------------------------------------
    title_box_ids = set()
    if 0 <= title_line_idx < len(lines):
        for box in lines[title_line_idx]:
            title_box_ids.add(id(box))

    header_box_ids = set()
    if title_line_idx > 0:
        for i in range(title_line_idx):
            for box in lines[i]:
                header_box_ids.add(id(box))

    # ------------------------------------------------------------------
    # 2. Vẽ từng OCR box (TITLE highlight đỏ)
    # ------------------------------------------------------------------
    for box in boxes:
        coords = box['coords']
        pts = np.array(coords, dtype=np.int32)
        if id(box) in title_box_ids:
            cv2.polylines(debug_img, [pts], True, (0, 0, 255), 3)
            overlay = debug_img.copy()
            cv2.fillPoly(overlay, [pts], (0, 0, 255))
            cv2.addWeighted(overlay, 0.15, debug_img, 0.85, 0, debug_img)
            _put_label_left(debug_img, "TITLE",
                            int(box['x_min']), int(box['y_min']), int(box['y_max']),
                            (0, 0, 255), LABEL_FONT_SCALE, LABEL_THICKNESS)

    # ------------------------------------------------------------------
    # 3. Vẽ bounding box HEADER (bao toàn bộ các dòng header)
    # ------------------------------------------------------------------
    if title_line_idx > 0:
        header_all_boxes = [b for line in lines[:title_line_idx] for b in line]
        if header_all_boxes:
            hx1 = int(min(b['x_min'] for b in header_all_boxes)) - 5
            hx2 = int(max(b['x_max'] for b in header_all_boxes)) + 5
            hy1 = int(min(b['y_min'] for b in header_all_boxes)) - 5
            hy2 = int(max(b['y_max'] for b in header_all_boxes)) + 5
            HEADER_COLOR = (180, 105, 255)   # hồng tím
            cv2.rectangle(debug_img, (hx1, hy1), (hx2, hy2), HEADER_COLOR, 2)
            overlay = debug_img.copy()
            cv2.rectangle(overlay, (hx1, hy1), (hx2, hy2), HEADER_COLOR, -1)
            cv2.addWeighted(overlay, 0.08, debug_img, 0.92, 0, debug_img)
            _put_label_left(debug_img, "HEADER",
                            hx1, hy1, hy2,
                            HEADER_COLOR, LABEL_FONT_SCALE, LABEL_THICKNESS)

    # ------------------------------------------------------------------
    # 4. Đường phân cách HEADER / BODY
    # ------------------------------------------------------------------
    if 0 <= title_line_idx < len(lines):
        title_line = lines[title_line_idx]
        split_y = int(max(b['y_max'] for b in title_line))
        cv2.line(debug_img, (0, split_y), (debug_img.shape[1], split_y), (255, 0, 0), 2)
        cv2.putText(debug_img, "--- HEADER / BODY SPLIT ---",
                    (10, split_y + 20),
                    cv2.FONT_HERSHEY_SIMPLEX, LABEL_FONT_SCALE, (255, 0, 0),
                    LABEL_THICKNESS, cv2.LINE_AA)

    # ------------------------------------------------------------------
    # 5. Vẽ bounding box EXT_TITLE
    # ------------------------------------------------------------------
    ext_start, ext_end = ext_title_range
    if ext_start < ext_end and ext_start < len(lines):
        ext_boxes = []
        for j in range(ext_start, min(ext_end, len(lines))):
            ext_boxes.extend(lines[j])
        if ext_boxes:
            ex1 = int(min(b['x_min'] for b in ext_boxes)) - 5
            ex2 = int(max(b['x_max'] for b in ext_boxes)) + 5
            ey1 = int(min(b['y_min'] for b in ext_boxes)) - 5
            ey2 = int(max(b['y_max'] for b in ext_boxes)) + 5
            EXT_COLOR = (0, 165, 255)
            cv2.rectangle(debug_img, (ex1, ey1), (ex2, ey2), EXT_COLOR, 3)
            overlay = debug_img.copy()
            cv2.rectangle(overlay, (ex1, ey1), (ex2, ey2), EXT_COLOR, -1)
            cv2.addWeighted(overlay, 0.12, debug_img, 0.88, 0, debug_img)
            _put_label_left(debug_img, "EXT_TITLE",
                            ex1, ey1, ey2,
                            EXT_COLOR, LABEL_FONT_SCALE, LABEL_THICKNESS)

    # ------------------------------------------------------------------
    # 6. Vẽ bounding box các đoạn BODY (P1, P2, ...)
    # ------------------------------------------------------------------
    if body_lines:
        line_stats, full_line_width, _ = compute_line_stats(body_lines, img_width)
        normal_gap = compute_normal_line_gap(line_stats)
        paragraphs = group_lines_into_paragraphs(line_stats, full_line_width, normal_gap)

        colors = [
            (0, 255, 0), (255, 165, 0), (255, 0, 255), (0, 255, 255),
            (255, 255, 0), (128, 0, 255), (0, 128, 255), (255, 0, 128),
        ]

        for para_idx, para_stats in enumerate(paragraphs):
            color = colors[para_idx % len(colors)]
            all_boxes_in_para = []
            for stat in para_stats:
                all_boxes_in_para.extend(stat['line_data'])
            if not all_boxes_in_para:
                continue

            px1 = int(min(b['x_min'] for b in all_boxes_in_para)) - 5
            px2 = int(max(b['x_max'] for b in all_boxes_in_para)) + 5
            py1 = int(min(b['y_min'] for b in all_boxes_in_para)) - 5
            py2 = int(max(b['y_max'] for b in all_boxes_in_para)) + 5

            cv2.rectangle(debug_img, (px1, py1), (px2, py2), color, 2)
            overlay = debug_img.copy()
            cv2.rectangle(overlay, (px1, py1), (px2, py2), color, -1)
            cv2.addWeighted(overlay, 0.08, debug_img, 0.92, 0, debug_img)
            _put_label_left(debug_img, f"P{para_idx+1}",
                            px1, py1, py2,
                            color, LABEL_FONT_SCALE, LABEL_THICKNESS)

    return debug_img


# ==========================================================================
# PHẦN 13: HÀM XỬ LÝ NHIỀU ẢNH (ENTRY POINT)
# ĐÃ SỬA: Bỏ apply_global_format
#          Enter: space_before=6pt, space_after=6pt
#          Title: line spacing 1.0, space_before=0, space_after=6pt
#          write_header_to_doc nhận thêm img_width
# ==========================================================================

def process_multiple_images(image_list, first_page_flags, progress_cb=None):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(18)

    debug_images = []
    all_info = []
    total_pages = len(image_list)

    for page_idx, (img_bgr, is_first) in enumerate(zip(image_list, first_page_flags)):
        page_label = "TRANG ĐẦU" if is_first else "TRANG THƯỜNG"
        print(f"\n{'='*60}")
        print(f"📄 TRANG {page_idx+1}/{total_pages} [{page_label}]")
        print(f"{'='*60}")

        page_start = page_idx / total_pages
        page_ocr_end = (page_idx + 0.6) / total_pages
        page_end = (page_idx + 1) / total_pages

        def ocr_progress(current, total, msg):
            if progress_cb and total > 0:
                frac = current / total
                overall = page_start + frac * (page_ocr_end - page_start)
                progress_cb(overall, 1.0, f"Trang {page_idx+1}: {msg}")

        boxes, avg_char_height, img_height, img_width = run_ocr_and_postprocess(
            img_bgr, progress_cb=ocr_progress
        )

        if not boxes:
            debug_images.append(img_bgr.copy())
            all_info.append({
                'total_boxes': 0, 'title_text': None,
                'title_line_idx': -1, 'ext_title_range': (0, 0),
                'num_header_lines': 0, 'num_body_lines': 0
            })
            continue

        lines = group_boxes_into_lines(boxes, avg_char_height)

        title_line_idx = -1
        title_text = None
        ext_title_range = (0, 0)

        if is_first:
            title_line_idx = find_title_line_index(lines, img_height, img_width)
            if title_line_idx >= 0:
                title_text = " ".join([b['text'] for b in lines[title_line_idx]]).strip()
                ext_title_range = find_extended_title_range(lines, title_line_idx)

        if page_idx > 0:
            doc.add_page_break()

        if is_first and title_line_idx >= 0:
            header_lines = lines[:title_line_idx]
            ext_start, ext_end = ext_title_range

            if header_lines:
                write_header_to_doc(doc, header_lines, avg_char_height, img_width)

            # Enter trước title: space_before=6pt, space_after=6pt
            empty_before = doc.add_paragraph()
            empty_before.paragraph_format.line_spacing = 1.0
            empty_before.paragraph_format.space_before = Pt(6)
            empty_before.paragraph_format.space_after = Pt(6)

            # Title chính: line spacing 1.0, space_before=0, space_after=6pt
            title_full = " ".join([b['text'] for b in lines[title_line_idx]]).strip()
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.line_spacing = 1.0
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after = Pt(6)
            run = p_title.add_run(title_full)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(15)
            run.bold = True

            # Extended title
            ext_title_lines = lines[ext_start:ext_end]
            if ext_title_lines:
                write_extended_title_to_doc(doc, ext_title_lines)

            # Enter sau ext_title: space_before=6pt, space_after=6pt
            empty_after = doc.add_paragraph()
            empty_after.paragraph_format.line_spacing = 1.0
            empty_after.paragraph_format.space_before = Pt(6)
            empty_after.paragraph_format.space_after = Pt(6)

            # Body
            body_start = ext_end
            remaining_body_lines = lines[body_start:]

            if remaining_body_lines:
                if progress_cb:
                    progress_cb(page_ocr_end, 1.0, f"Trang {page_idx+1}: Xuất body...")
                write_body_to_doc(doc, remaining_body_lines, avg_char_height, img_width)

            debug_img = draw_debug_image(
                img_bgr, boxes, lines, title_line_idx,
                ext_title_range, remaining_body_lines,
                avg_char_height, img_width
            )

            all_info.append({
                'total_boxes': len(boxes), 'title_text': title_text,
                'title_line_idx': title_line_idx,
                'ext_title_range': ext_title_range,
                'num_header_lines': len(header_lines),
                'num_body_lines': len(remaining_body_lines) if remaining_body_lines else 0,
            })

        else:
            if progress_cb:
                progress_cb(page_ocr_end, 1.0, f"Trang {page_idx+1}: Xuất body...")
            write_body_to_doc(doc, lines, avg_char_height, img_width)

            debug_img = draw_debug_image(
                img_bgr, boxes, lines, -1,
                (0, 0), lines,
                avg_char_height, img_width
            )

            all_info.append({
                'total_boxes': len(boxes), 'title_text': None,
                'title_line_idx': -1, 'ext_title_range': (0, 0),
                'num_header_lines': 0,
                'num_body_lines': len(lines),
            })

        debug_images.append(debug_img)

        if progress_cb:
            progress_cb(page_end, 1.0, f"Trang {page_idx+1}: Hoàn tất!")

    return doc, debug_images, all_info
